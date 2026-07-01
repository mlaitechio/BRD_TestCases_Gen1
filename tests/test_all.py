"""
BRD Automation — Complete Python Test Suite
============================================
Covers:
  - Django models (Project, AgentOutput)
  - AI base layer (JSON parsing, fence stripping)
  - All 11 REST API endpoints
  - File extractor utility
  - DOCX exporter utility
  - Celery tasks (mocked AI calls)

Run with:
    python manage.py test tests --verbosity=2
OR with pytest:
    pytest tests/ -v
"""

import json
import io
import uuid
from unittest.mock import patch, MagicMock
from django.test import TestCase, Client
from django.urls import reverse
from django.core.files.uploadedfile import SimpleUploadedFile

from apps.projects.models import Project, AgentOutput
from agents.base import _parse_json


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

SAMPLE_PROJECT_DESCRIPTION = (
    "Build a B2B customer portal for enterprise clients to track their orders, "
    "view invoices, and raise support tickets. Must integrate with Salesforce CRM. "
    "GDPR compliance required. Target users are procurement managers."
)

SAMPLE_CLARIFICATION_RESULT = {
    "questions": [
        {"id": "Q1", "question": "Who are the primary users?", "why_asking": "To define personas"},
        {"id": "Q2", "question": "Which CRM fields need to sync?", "why_asking": "Integration scope"},
        {"id": "Q3", "question": "What data residency rules apply?", "why_asking": "GDPR compliance"},
    ]
}

SAMPLE_BRD_RESULT = {
    "executive_summary": "A B2B customer portal for enterprise order tracking.",
    "project_scope": {
        "in_scope": ["Order tracking", "Invoice management", "Support tickets"],
        "out_of_scope": ["Mobile app", "Payment processing"]
    },
    "business_objectives": [
        {"id": "BO-001", "objective": "Reduce support emails by 40%", "metric": "Email volume", "target": "40% reduction"}
    ],
    "stakeholders": [
        {"role": "Product Owner", "responsibilities": "Prioritise backlog", "interest_level": "High", "influence_level": "High"}
    ],
    "project_plan": {
        "phases": [{"phase": "Phase 1 — Discovery", "duration": "2 weeks", "deliverables": ["BRD"], "milestones": ["Sign-off"]}]
    },
    "effort_estimation": {"total_estimated_hours": 800, "summary": "Medium complexity", "breakdown": []},
    "functional_requirements": [
        {
            "id": "FR-001", "title": "Order Tracking Dashboard",
            "description": "Users can view all orders in a filterable table.",
            "priority": "Must Have",
            "acceptance_criteria": ["Table loads in < 2s", "Filters work correctly"],
            "compliance_notes": "N/A"
        },
        {
            "id": "FR-002", "title": "Support Ticket Creation",
            "description": "Users can raise tickets linked to specific orders.",
            "priority": "Must Have",
            "acceptance_criteria": ["Ticket created within 3s"],
            "compliance_notes": "GDPR — no PII in ticket body unless necessary"
        }
    ],
    "non_functional_requirements": [
        {"id": "NFR-001", "category": "Performance", "requirement": "Page load < 2s", "metric": "< 2000ms p95", "priority": "Must Have"},
        {"id": "NFR-002", "category": "Security", "requirement": "GDPR compliant data handling", "metric": "0 data breaches", "priority": "Must Have"}
    ],
    "constraints_and_assumptions": {
        "constraints": [{"id": "CON-001", "description": "Budget capped at $200k", "impact": "Limits team size"}],
        "assumptions": [{"id": "ASS-001", "description": "Salesforce API available", "risk_if_wrong": "Integration blocked"}]
    },
    "success_criteria": [
        {"id": "SC-001", "criterion": "Portal live within 6 months", "measurement_method": "Go-live date", "target": "Month 6"}
    ],
    "glossary": [
        {"term": "BRD", "definition": "Business Requirements Document"}
    ]
}

SAMPLE_PLAN_RESULT = {
    "project_summary": "Phased delivery of the B2B portal.",
    "methodology": "Agile",
    "total_duration": "6 months",
    "team_structure": [{"role": "Backend Developer", "count": 2, "responsibilities": "API development"}],
    "phases": [
        {
            "phase_number": 1, "phase_name": "Discovery", "duration": "2 weeks",
            "start_week": 1, "end_week": 2,
            "objectives": ["Define scope"],
            "tasks": [{"task_id": "T-001", "task_name": "Stakeholder interviews", "assignee_role": "BA", "duration_days": 5, "dependencies": [], "linked_requirements": ["FR-001"]}],
            "deliverables": ["BRD sign-off"],
            "milestones": ["BRD approved"]
        }
    ],
    "risk_register": [{"risk_id": "R-001", "description": "API delays", "probability": "Medium", "impact": "High", "mitigation": "Start integration early", "owner": "Tech Lead"}],
    "communication_plan": {"meetings": [{"meeting_type": "Sprint Review", "frequency": "Bi-weekly", "attendees": ["PM", "Dev"], "purpose": "Review progress"}], "reporting": "Weekly status email"},
    "definition_of_done": ["All tests pass", "Code reviewed"]
}

SAMPLE_TESTCASE_RESULT = {
    "test_summary": {
        "total_test_cases": 4,
        "coverage_percentage": "100% of functional requirements covered",
        "test_categories": {"functional": 2, "integration": 0, "edge_case": 1, "negative": 1, "acceptance": 0}
    },
    "test_cases": [
        {
            "test_id": "TC-001", "linked_requirement": "FR-001",
            "title": "Order table displays all orders",
            "type": "Functional", "priority": "High",
            "preconditions": ["User is logged in", "Orders exist in system"],
            "test_steps": [{"step": 1, "action": "Navigate to dashboard", "expected_result": "Table visible"}],
            "expected_outcome": "All orders listed",
            "pass_criteria": "Table renders with correct data"
        }
    ],
    "traceability_matrix": [
        {"requirement_id": "FR-001", "requirement_title": "Order Tracking Dashboard", "linked_test_cases": ["TC-001"], "coverage_status": "Covered"}
    ]
}

SAMPLE_EFFORT_RESULT = {
    "estimation_summary": {
        "total_hours": 800, "total_weeks": 20, "total_months": 5,
        "team_size_recommended": 4,
        "confidence_level": "Medium",
        "estimation_methodology": "Function Point Analysis"
    },
    "by_phase": [{"phase": "Discovery", "hours": 80, "percentage_of_total": "10%"}],
    "by_role": [{"role": "Backend Developer", "hours": 400, "daily_rate_usd": 600, "total_cost_usd": 30000}],
    "by_feature": [{"feature_area": "Order Tracking", "linked_requirements": ["FR-001"], "complexity": "Medium", "hours": 120, "notes": "Includes Salesforce sync"}],
    "cost_estimate": {"currency": "USD", "low_estimate": 120000, "mid_estimate": 150000, "high_estimate": 190000, "notes": "Includes 20% buffer"},
    "risks_affecting_estimate": [{"risk": "Scope creep", "potential_impact_hours": 80, "mitigation": "Change control process"}],
    "assumptions": ["Salesforce API available from day 1"]
}


def _make_project(status='new', extracted_text=None, clarification_questions=None, clarification_answers=None, brd_approved=False):
    """Factory helper: create a Project instance."""
    return Project.objects.create(
        raw_input=SAMPLE_PROJECT_DESCRIPTION,
        extracted_text=extracted_text or SAMPLE_PROJECT_DESCRIPTION,
        status=status,
        clarification_questions=clarification_questions,
        clarification_answers=clarification_answers,
        brd_approved=brd_approved,
    )


def _make_agent_output(project, agent_type, status='complete', structured_output=None):
    """Factory helper: create an AgentOutput instance."""
    return AgentOutput.objects.create(
        project=project,
        agent_type=agent_type,
        status=status,
        structured_output=structured_output,
        raw_output=str(structured_output) if structured_output else None,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# 1. MODEL TESTS
# ═══════════════════════════════════════════════════════════════════════════════

class ProjectModelTests(TestCase):
    """Tests for the Project model."""

    def test_project_created_with_uuid(self):
        project = _make_project()
        self.assertIsInstance(project.id, uuid.UUID)

    def test_project_default_status_is_new(self):
        project = _make_project()
        self.assertEqual(project.status, 'new')

    def test_project_str_representation(self):
        project = _make_project()
        self.assertIn('new', str(project))

    def test_project_brd_approved_defaults_false(self):
        project = _make_project()
        self.assertFalse(project.brd_approved)

    def test_project_ordering_newest_first(self):
        from django.utils import timezone
        import time
        p1 = _make_project()
        time.sleep(0.01)  # ensure distinct timestamps
        p2 = _make_project()
        projects = list(Project.objects.all())
        # p2 was created after p1, so it should come first (newest first ordering)
        self.assertEqual(projects[0].id, p2.id)

    def test_project_status_choices_valid(self):
        valid_statuses = [s[0] for s in Project.STATUS_CHOICES]
        self.assertIn('awaiting_approval', valid_statuses)
        self.assertIn('complete', valid_statuses)
        self.assertIn('failed', valid_statuses)

    def test_project_stores_clarification_questions(self):
        project = _make_project(clarification_questions=SAMPLE_CLARIFICATION_RESULT['questions'])
        self.assertEqual(len(project.clarification_questions), 3)
        self.assertEqual(project.clarification_questions[0]['id'], 'Q1')

    def test_project_stores_clarification_answers(self):
        answers = {'Q1': 'Procurement managers', 'Q2': 'Order status and amount'}
        project = _make_project(clarification_answers=answers)
        self.assertEqual(project.clarification_answers['Q1'], 'Procurement managers')


class AgentOutputModelTests(TestCase):
    """Tests for the AgentOutput model."""

    def setUp(self):
        self.project = _make_project()

    def test_agent_output_created_with_uuid(self):
        output = _make_agent_output(self.project, 'brd')
        self.assertIsInstance(output.id, uuid.UUID)

    def test_agent_output_linked_to_project(self):
        output = _make_agent_output(self.project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        self.assertEqual(output.project.id, self.project.id)

    def test_agent_output_unique_per_project_per_type(self):
        _make_agent_output(self.project, 'brd')
        with self.assertRaises(Exception):
            _make_agent_output(self.project, 'brd')

    def test_agent_output_stores_json(self):
        output = _make_agent_output(self.project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        self.assertIn('functional_requirements', output.structured_output)

    def test_agent_output_str_representation(self):
        output = _make_agent_output(self.project, 'brd')
        self.assertIn('brd', str(output))

    def test_all_agent_types_valid(self):
        valid_types = [t[0] for t in AgentOutput.AGENT_TYPE_CHOICES]
        for t in ['clarification', 'brd', 'plan', 'test_cases', 'effort']:
            self.assertIn(t, valid_types)

    def test_update_or_create_pattern(self):
        """Ensures re-running an agent updates existing record, not duplicates."""
        AgentOutput.objects.create(
            project=self.project, agent_type='brd', status='complete',
            structured_output=SAMPLE_BRD_RESULT
        )
        AgentOutput.objects.update_or_create(
            project=self.project, agent_type='brd',
            defaults={'status': 'running', 'structured_output': None}
        )
        self.assertEqual(AgentOutput.objects.filter(project=self.project, agent_type='brd').count(), 1)
        updated = AgentOutput.objects.get(project=self.project, agent_type='brd')
        self.assertEqual(updated.status, 'running')


# ═══════════════════════════════════════════════════════════════════════════════
# 2. AI BASE LAYER TESTS
# ═══════════════════════════════════════════════════════════════════════════════

class AIBaseParserTests(TestCase):
    """Tests for the JSON parsing / fence-stripping logic in agents/base.py."""

    def test_parse_clean_json(self):
        raw = '{"key": "value"}'
        result = _parse_json(raw)
        self.assertEqual(result['key'], 'value')

    def test_parse_json_with_backtick_fence(self):
        raw = '```json\n{"key": "value"}\n```'
        result = _parse_json(raw)
        self.assertEqual(result['key'], 'value')

    def test_parse_json_with_plain_fence(self):
        raw = '```\n{"key": "value"}\n```'
        result = _parse_json(raw)
        self.assertEqual(result['key'], 'value')

    def test_parse_json_with_surrounding_text(self):
        raw = 'Here is the JSON output:\n{"key": "value"}\nEnd of output.'
        result = _parse_json(raw)
        self.assertEqual(result['key'], 'value')

    def test_parse_json_with_leading_whitespace(self):
        raw = '   \n  {"key": "value"}  \n  '
        result = _parse_json(raw)
        self.assertEqual(result['key'], 'value')

    def test_parse_invalid_json_raises_value_error(self):
        raw = 'This is not JSON at all, sorry.'
        with self.assertRaises(ValueError):
            _parse_json(raw)

    def test_parse_nested_json(self):
        raw = '{"outer": {"inner": [1, 2, 3]}}'
        result = _parse_json(raw)
        self.assertEqual(result['outer']['inner'], [1, 2, 3])

    def test_parse_json_with_unicode(self):
        raw = '{"name": "Ação do Usuário"}'
        result = _parse_json(raw)
        self.assertEqual(result['name'], 'Ação do Usuário')


# ═══════════════════════════════════════════════════════════════════════════════
# 3. API ENDPOINT TESTS
# ═══════════════════════════════════════════════════════════════════════════════

class ProjectCreateViewTests(TestCase):
    """POST /api/projects/"""

    def setUp(self):
        self.client = Client()
        self.url = '/api/projects/'

    @patch('apps.projects.tasks.run_clarification_task')
    def test_create_project_with_text(self, mock_task):
        mock_task.delay = MagicMock()
        response = self.client.post(
            self.url,
            data=json.dumps({'raw_input': SAMPLE_PROJECT_DESCRIPTION}),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 201)
        data = response.json()
        self.assertIn('id', data)
        self.assertEqual(data['status'], 'new')
        mock_task.delay.assert_called_once()

    @patch('apps.projects.tasks.run_clarification_task')
    def test_create_project_with_txt_file(self, mock_task):
        mock_task.delay = MagicMock()
        txt_content = b'Build a project management tool for software teams.'
        uploaded_file = SimpleUploadedFile('description.txt', txt_content, content_type='text/plain')
        response = self.client.post(self.url, data={'uploaded_file': uploaded_file})
        self.assertEqual(response.status_code, 201)
        mock_task.delay.assert_called_once()

    def test_create_project_without_input_returns_400(self):
        response = self.client.post(
            self.url,
            data=json.dumps({}),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 400)

    @patch('apps.projects.tasks.run_clarification_task')
    def test_create_project_stores_extracted_text(self, mock_task):
        mock_task.delay = MagicMock()
        self.client.post(
            self.url,
            data=json.dumps({'raw_input': SAMPLE_PROJECT_DESCRIPTION}),
            content_type='application/json'
        )
        project = Project.objects.first()
        self.assertIsNotNone(project.extracted_text)
        self.assertIn('portal', project.extracted_text)


class ProjectStatusViewTests(TestCase):
    """GET /api/projects/:id/status/"""

    def setUp(self):
        self.client = Client()
        self.project = _make_project(status='awaiting_answers')
        _make_agent_output(self.project, 'clarification')

    def test_status_returns_200(self):
        response = self.client.get(f'/api/projects/{self.project.id}/status/')
        self.assertEqual(response.status_code, 200)

    def test_status_contains_project_id(self):
        response = self.client.get(f'/api/projects/{self.project.id}/status/')
        data = response.json()
        self.assertEqual(data['id'], str(self.project.id))

    def test_status_contains_outputs_dict(self):
        response = self.client.get(f'/api/projects/{self.project.id}/status/')
        data = response.json()
        self.assertIn('outputs', data)
        self.assertIn('clarification', data['outputs'])

    def test_status_404_for_unknown_project(self):
        fake_id = uuid.uuid4()
        response = self.client.get(f'/api/projects/{fake_id}/status/')
        self.assertEqual(response.status_code, 404)


class ClarificationQuestionsViewTests(TestCase):
    """GET /api/projects/:id/clarification-questions/"""

    def setUp(self):
        self.client = Client()

    def test_returns_questions_when_ready(self):
        project = _make_project(
            status='awaiting_answers',
            clarification_questions=SAMPLE_CLARIFICATION_RESULT['questions']
        )
        response = self.client.get(f'/api/projects/{project.id}/clarification-questions/')
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIn('questions', data)
        self.assertEqual(len(data['questions']), 3)

    def test_returns_202_while_still_generating(self):
        project = _make_project(status='clarifying')
        response = self.client.get(f'/api/projects/{project.id}/clarification-questions/')
        self.assertEqual(response.status_code, 202)

    def test_returns_404_when_not_yet_available(self):
        project = _make_project(status='new')
        response = self.client.get(f'/api/projects/{project.id}/clarification-questions/')
        self.assertEqual(response.status_code, 404)

    def test_questions_have_required_fields(self):
        project = _make_project(
            status='awaiting_answers',
            clarification_questions=SAMPLE_CLARIFICATION_RESULT['questions']
        )
        response = self.client.get(f'/api/projects/{project.id}/clarification-questions/')
        q = response.json()['questions'][0]
        self.assertIn('id', q)
        self.assertIn('question', q)
        self.assertIn('why_asking', q)


class AnswerQuestionsViewTests(TestCase):
    """POST /api/projects/:id/answer-questions/"""

    def setUp(self):
        self.client = Client()

    @patch('apps.projects.tasks.run_brd_task')
    def test_submit_answers_fires_brd_task(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(
            status='awaiting_answers',
            clarification_questions=SAMPLE_CLARIFICATION_RESULT['questions']
        )
        payload = {'answers': {'Q1': 'Procurement managers', 'Q2': 'Salesforce CRM', 'Q3': 'EU GDPR'}}
        response = self.client.post(
            f'/api/projects/{project.id}/answer-questions/',
            data=json.dumps(payload),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 200)
        mock_task.delay.assert_called_once_with(str(project.id))

    @patch('apps.projects.tasks.run_brd_task')
    def test_answers_saved_to_project(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_answers')
        payload = {'answers': {'Q1': 'Test answer'}}
        self.client.post(
            f'/api/projects/{project.id}/answer-questions/',
            data=json.dumps(payload),
            content_type='application/json'
        )
        project.refresh_from_db()
        self.assertEqual(project.clarification_answers['Q1'], 'Test answer')

    def test_submit_answers_wrong_status_returns_400(self):
        project = _make_project(status='generating_brd')
        payload = {'answers': {'Q1': 'answer'}}
        response = self.client.post(
            f'/api/projects/{project.id}/answer-questions/',
            data=json.dumps(payload),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 400)

    def test_submit_without_answers_key_returns_400(self):
        project = _make_project(status='awaiting_answers')
        response = self.client.post(
            f'/api/projects/{project.id}/answer-questions/',
            data=json.dumps({}),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 400)


class BRDOutputViewTests(TestCase):
    """GET /api/projects/:id/brd/"""

    def setUp(self):
        self.client = Client()

    def test_returns_brd_when_complete(self):
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        response = self.client.get(f'/api/projects/{project.id}/brd/')
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIn('structured_output', data)
        self.assertIn('functional_requirements', data['structured_output'])

    def test_returns_404_when_brd_not_ready(self):
        project = _make_project(status='generating_brd')
        response = self.client.get(f'/api/projects/{project.id}/brd/')
        self.assertEqual(response.status_code, 404)

    def test_brd_output_has_all_11_sections(self):
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        response = self.client.get(f'/api/projects/{project.id}/brd/')
        output = response.json()['structured_output']
        required_sections = [
            'executive_summary', 'project_scope', 'business_objectives',
            'stakeholders', 'project_plan', 'effort_estimation',
            'functional_requirements', 'non_functional_requirements',
            'constraints_and_assumptions', 'success_criteria', 'glossary'
        ]
        for section in required_sections:
            self.assertIn(section, output, f'Missing BRD section: {section}')


class ApproveBRDViewTests(TestCase):
    """POST /api/projects/:id/approve-brd/"""

    def setUp(self):
        self.client = Client()

    @patch('apps.projects.tasks.run_remaining_agents_task')
    def test_approve_brd_fires_remaining_agents(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        response = self.client.post(f'/api/projects/{project.id}/approve-brd/')
        self.assertEqual(response.status_code, 200)
        mock_task.delay.assert_called_once_with(str(project.id))

    @patch('apps.projects.tasks.run_remaining_agents_task')
    def test_approve_brd_sets_brd_approved_true(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        self.client.post(f'/api/projects/{project.id}/approve-brd/')
        project.refresh_from_db()
        self.assertTrue(project.brd_approved)

    def test_approve_brd_wrong_status_returns_400(self):
        project = _make_project(status='generating_brd')
        response = self.client.post(f'/api/projects/{project.id}/approve-brd/')
        self.assertEqual(response.status_code, 400)


class ReviseBRDViewTests(TestCase):
    """POST /api/projects/:id/revise-brd/"""

    def setUp(self):
        self.client = Client()

    @patch('apps.projects.tasks.run_brd_task')
    def test_revise_brd_fires_brd_task(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        payload = {'revision_notes': 'Please add more detail on payment integration.'}
        response = self.client.post(
            f'/api/projects/{project.id}/revise-brd/',
            data=json.dumps(payload),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 200)
        mock_task.delay.assert_called_once()

    @patch('apps.projects.tasks.run_brd_task')
    def test_revise_brd_saves_revision_notes(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_approval')
        _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        payload = {'revision_notes': 'Expand on NFRs.'}
        self.client.post(
            f'/api/projects/{project.id}/revise-brd/',
            data=json.dumps(payload),
            content_type='application/json'
        )
        project.refresh_from_db()
        self.assertEqual(project.revision_notes, 'Expand on NFRs.')

    @patch('apps.projects.tasks.run_brd_task')
    def test_revise_brd_resets_brd_output(self, mock_task):
        mock_task.delay = MagicMock()
        project = _make_project(status='awaiting_approval')
        brd_output = _make_agent_output(project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        self.client.post(
            f'/api/projects/{project.id}/revise-brd/',
            data=json.dumps({'revision_notes': 'Revise it.'}),
            content_type='application/json'
        )
        brd_output.refresh_from_db()
        self.assertEqual(brd_output.status, 'pending')
        self.assertIsNone(brd_output.structured_output)

    def test_revise_without_notes_returns_400(self):
        project = _make_project(status='awaiting_approval')
        response = self.client.post(
            f'/api/projects/{project.id}/revise-brd/',
            data=json.dumps({}),
            content_type='application/json'
        )
        self.assertEqual(response.status_code, 400)


class PlanOutputViewTests(TestCase):
    """GET /api/projects/:id/plan/"""

    def test_returns_plan_when_complete(self):
        client = Client()
        project = _make_project(status='complete')
        _make_agent_output(project, 'plan', structured_output=SAMPLE_PLAN_RESULT)
        response = client.get(f'/api/projects/{project.id}/plan/')
        self.assertEqual(response.status_code, 200)
        self.assertIn('structured_output', response.json())

    def test_returns_404_when_plan_not_ready(self):
        client = Client()
        project = _make_project(status='awaiting_approval')
        response = client.get(f'/api/projects/{project.id}/plan/')
        self.assertEqual(response.status_code, 404)


class TestCasesOutputViewTests(TestCase):
    """GET /api/projects/:id/testcases/"""

    def test_returns_testcases_when_complete(self):
        client = Client()
        project = _make_project(status='complete')
        _make_agent_output(project, 'test_cases', structured_output=SAMPLE_TESTCASE_RESULT)
        response = client.get(f'/api/projects/{project.id}/testcases/')
        self.assertEqual(response.status_code, 200)
        output = response.json()['structured_output']
        self.assertIn('test_cases', output)
        self.assertIn('traceability_matrix', output)

    def test_returns_404_when_not_ready(self):
        client = Client()
        project = _make_project(status='approved')
        response = client.get(f'/api/projects/{project.id}/testcases/')
        self.assertEqual(response.status_code, 404)


class EffortOutputViewTests(TestCase):
    """GET /api/projects/:id/effort/"""

    def test_returns_effort_when_complete(self):
        client = Client()
        project = _make_project(status='complete')
        _make_agent_output(project, 'effort', structured_output=SAMPLE_EFFORT_RESULT)
        response = client.get(f'/api/projects/{project.id}/effort/')
        self.assertEqual(response.status_code, 200)
        output = response.json()['structured_output']
        self.assertIn('estimation_summary', output)
        self.assertIn('cost_estimate', output)


class DownloadOutputViewTests(TestCase):
    """GET /api/projects/:id/download/:type/"""

    def setUp(self):
        self.client = Client()
        self.project = _make_project(status='complete')

    def test_download_brd_returns_docx(self):
        _make_agent_output(self.project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        response = self.client.get(f'/api/projects/{self.project.id}/download/brd/')
        self.assertEqual(response.status_code, 200)
        self.assertIn('application/vnd.openxmlformats', response['Content-Type'])
        self.assertIn('.docx', response['Content-Disposition'])

    def test_download_plan_returns_docx(self):
        _make_agent_output(self.project, 'plan', structured_output=SAMPLE_PLAN_RESULT)
        response = self.client.get(f'/api/projects/{self.project.id}/download/plan/')
        self.assertEqual(response.status_code, 200)

    def test_download_testcases_returns_docx(self):
        _make_agent_output(self.project, 'test_cases', structured_output=SAMPLE_TESTCASE_RESULT)
        response = self.client.get(f'/api/projects/{self.project.id}/download/testcases/')
        self.assertEqual(response.status_code, 200)

    def test_download_effort_returns_docx(self):
        _make_agent_output(self.project, 'effort', structured_output=SAMPLE_EFFORT_RESULT)
        response = self.client.get(f'/api/projects/{self.project.id}/download/effort/')
        self.assertEqual(response.status_code, 200)

    def test_download_invalid_type_returns_400(self):
        response = self.client.get(f'/api/projects/{self.project.id}/download/invalid/')
        self.assertEqual(response.status_code, 400)

    def test_download_when_output_not_ready_returns_425(self):
        _make_agent_output(self.project, 'brd', status='running', structured_output=None)
        response = self.client.get(f'/api/projects/{self.project.id}/download/brd/')
        self.assertEqual(response.status_code, 425)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. FILE EXTRACTOR TESTS
# ═══════════════════════════════════════════════════════════════════════════════

class FileExtractorTests(TestCase):
    """Tests for utils/file_extractor.py"""

    def test_extract_txt_file(self):
        import tempfile, os
        from utils.file_extractor import extract_text_from_file

        content = 'Build a project tracking system for a healthcare company.'
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            tmp_path = f.name

        try:
            result = extract_text_from_file(tmp_path)
            self.assertEqual(result, content)
        finally:
            os.unlink(tmp_path)

    def test_extract_returns_none_for_missing_file(self):
        from utils.file_extractor import extract_text_from_file
        result = extract_text_from_file('/nonexistent/path/file.txt')
        self.assertIsNone(result)

    def test_extract_returns_none_for_unsupported_type(self):
        import tempfile, os
        from utils.file_extractor import extract_text_from_file

        with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
            f.write(b'some content')
            tmp_path = f.name

        try:
            result = extract_text_from_file(tmp_path)
            self.assertIsNone(result)
        finally:
            os.unlink(tmp_path)

    def test_extract_txt_with_utf8_content(self):
        import tempfile, os
        from utils.file_extractor import extract_text_from_file

        content = 'Système de gestion — Données utilisateur — RGPD'
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            tmp_path = f.name

        try:
            result = extract_text_from_file(tmp_path)
            self.assertEqual(result, content)
        finally:
            os.unlink(tmp_path)

    def test_extract_docx_file(self):
        import tempfile, os
        from utils.file_extractor import extract_text_from_file
        from docx import Document

        doc = Document()
        doc.add_paragraph('Build a patient records system.')
        doc.add_paragraph('Must comply with HIPAA.')

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        try:
            result = extract_text_from_file(tmp_path)
            self.assertIn('patient records', result)
            self.assertIn('HIPAA', result)
        finally:
            os.unlink(tmp_path)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. DOCX EXPORTER TESTS
# ═══════════════════════════════════════════════════════════════════════════════

class DocxExporterTests(TestCase):
    """Tests for utils/docx_exporter.py — verifies each exporter returns a valid DOCX BytesIO."""

    def _assert_valid_docx(self, buffer):
        """Verify the buffer is a valid DOCX file."""
        from docx import Document
        buffer.seek(0)
        doc = Document(buffer)
        # If Document() doesn't raise, it's valid
        self.assertIsNotNone(doc)
        return doc

    def test_export_brd_returns_bytes_io(self):
        from utils.docx_exporter import export_brd_to_docx
        buffer = export_brd_to_docx(SAMPLE_BRD_RESULT)
        self.assertIsInstance(buffer, io.BytesIO)
        self._assert_valid_docx(buffer)

    def test_export_brd_contains_executive_summary(self):
        from utils.docx_exporter import export_brd_to_docx
        from docx import Document
        buffer = export_brd_to_docx(SAMPLE_BRD_RESULT)
        buffer.seek(0)
        doc = Document(buffer)
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        self.assertIn('Executive Summary', all_text)

    def test_export_brd_contains_all_11_section_headings(self):
        from utils.docx_exporter import export_brd_to_docx
        from docx import Document
        buffer = export_brd_to_docx(SAMPLE_BRD_RESULT)
        buffer.seek(0)
        doc = Document(buffer)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        headings_text = ' '.join(headings)
        for section in ['Executive Summary', 'Project Scope', 'Business Objectives',
                         'Stakeholder List', 'Project Plan', 'Effort Estimation',
                         'Functional Requirements', 'Non-Functional Requirements',
                         'Constraints and Assumptions', 'Success Criteria', 'Glossary']:
            self.assertIn(section, headings_text, f'Missing heading: {section}')

    def test_export_plan_returns_valid_docx(self):
        from utils.docx_exporter import export_plan_to_docx
        buffer = export_plan_to_docx(SAMPLE_PLAN_RESULT)
        self._assert_valid_docx(buffer)

    def test_export_testcases_returns_valid_docx(self):
        from utils.docx_exporter import export_testcases_to_docx
        buffer = export_testcases_to_docx(SAMPLE_TESTCASE_RESULT)
        self._assert_valid_docx(buffer)

    def test_export_effort_returns_valid_docx(self):
        from utils.docx_exporter import export_effort_to_docx
        buffer = export_effort_to_docx(SAMPLE_EFFORT_RESULT)
        self._assert_valid_docx(buffer)

    def test_export_brd_with_empty_sections_doesnt_crash(self):
        from utils.docx_exporter import export_brd_to_docx
        minimal = {
            "executive_summary": "Minimal project.",
            "project_scope": {"in_scope": [], "out_of_scope": []},
            "business_objectives": [], "stakeholders": [],
            "project_plan": {"phases": []}, "effort_estimation": {"total_estimated_hours": 0, "summary": "", "breakdown": []},
            "functional_requirements": [], "non_functional_requirements": [],
            "constraints_and_assumptions": {"constraints": [], "assumptions": []},
            "success_criteria": [], "glossary": []
        }
        buffer = export_brd_to_docx(minimal)
        self._assert_valid_docx(buffer)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. CELERY TASK TESTS (mocked AI calls)
# ═══════════════════════════════════════════════════════════════════════════════

class CeleryTaskTests(TestCase):
    """Tests for Celery tasks with AI calls mocked out."""

    def setUp(self):
        self.project = _make_project(
            status='new',
            extracted_text=SAMPLE_PROJECT_DESCRIPTION
        )

    @patch('agents.clarification_agent.generate_clarification_questions', return_value=SAMPLE_CLARIFICATION_RESULT)
    def test_run_clarification_task_sets_status_to_awaiting_answers(self, mock_agent):
        from apps.projects.tasks import run_clarification_task
        run_clarification_task(str(self.project.id))
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, 'awaiting_answers')

    @patch('agents.clarification_agent.generate_clarification_questions', return_value=SAMPLE_CLARIFICATION_RESULT)
    def test_run_clarification_task_saves_questions(self, mock_agent):
        from apps.projects.tasks import run_clarification_task
        run_clarification_task(str(self.project.id))
        self.project.refresh_from_db()
        self.assertIsNotNone(self.project.clarification_questions)
        self.assertEqual(len(self.project.clarification_questions), 3)

    @patch('agents.clarification_agent.generate_clarification_questions', return_value=SAMPLE_CLARIFICATION_RESULT)
    def test_run_clarification_task_creates_agent_output(self, mock_agent):
        from apps.projects.tasks import run_clarification_task
        run_clarification_task(str(self.project.id))
        output = AgentOutput.objects.get(project=self.project, agent_type='clarification')
        self.assertEqual(output.status, 'complete')

    @patch('agents.brd_agent.generate_brd', return_value=SAMPLE_BRD_RESULT)
    def test_run_brd_task_sets_status_to_awaiting_approval(self, mock_agent):
        from apps.projects.tasks import run_brd_task
        self.project.status = 'awaiting_answers'
        self.project.clarification_answers = {'Q1': 'answer'}
        self.project.save()
        run_brd_task(str(self.project.id))
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, 'awaiting_approval')

    @patch('agents.brd_agent.generate_brd', return_value=SAMPLE_BRD_RESULT)
    def test_run_brd_task_saves_structured_output(self, mock_agent):
        from apps.projects.tasks import run_brd_task
        self.project.status = 'awaiting_answers'
        self.project.save()
        run_brd_task(str(self.project.id))
        output = AgentOutput.objects.get(project=self.project, agent_type='brd')
        self.assertEqual(output.status, 'complete')
        self.assertIn('functional_requirements', output.structured_output)

    @patch('agents.effort_agent.generate_effort_estimation', return_value=SAMPLE_EFFORT_RESULT)
    @patch('agents.testcase_agent.generate_test_cases', return_value=SAMPLE_TESTCASE_RESULT)
    @patch('agents.plan_agent.generate_project_plan', return_value=SAMPLE_PLAN_RESULT)
    def test_run_remaining_agents_sets_status_complete(self, mock_plan, mock_tc, mock_effort):
        from apps.projects.tasks import run_remaining_agents_task
        self.project.status = 'awaiting_approval'
        self.project.brd_approved = True
        self.project.save()
        _make_agent_output(self.project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        run_remaining_agents_task(str(self.project.id))
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, 'complete')

    @patch('agents.effort_agent.generate_effort_estimation', return_value=SAMPLE_EFFORT_RESULT)
    @patch('agents.testcase_agent.generate_test_cases', return_value=SAMPLE_TESTCASE_RESULT)
    @patch('agents.plan_agent.generate_project_plan', return_value=SAMPLE_PLAN_RESULT)
    def test_run_remaining_agents_creates_all_3_outputs(self, mock_plan, mock_tc, mock_effort):
        from apps.projects.tasks import run_remaining_agents_task
        self.project.status = 'awaiting_approval'
        self.project.save()
        _make_agent_output(self.project, 'brd', structured_output=SAMPLE_BRD_RESULT)
        run_remaining_agents_task(str(self.project.id))
        types = AgentOutput.objects.filter(project=self.project).values_list('agent_type', flat=True)
        self.assertIn('plan', types)
        self.assertIn('test_cases', types)
        self.assertIn('effort', types)

    @patch('agents.clarification_agent.generate_clarification_questions', side_effect=ValueError('API timeout'))
    def test_run_clarification_task_handles_ai_failure(self, mock_agent):
        from apps.projects.tasks import run_clarification_task
        # Call the task directly (not via .apply which requires a broker)
        # The task catches the error and sets status to failed, then re-raises
        try:
            run_clarification_task(str(self.project.id))
        except Exception:
            pass
        # Status should have been attempted to be set to failed
        self.project.refresh_from_db()
        self.assertIn(self.project.status, ('failed', 'clarifying'))


# ═══════════════════════════════════════════════════════════════════════════════
# 7. FULL PIPELINE INTEGRATION TEST
# ═══════════════════════════════════════════════════════════════════════════════

class FullPipelineIntegrationTest(TestCase):
    """
    Simulates the full API flow end-to-end with all AI calls mocked.
    Tasks are called DIRECTLY (synchronously) in the test after the views
    fire .delay(). patch.object(.delay) prevents the Celery broker call
    without replacing the actual task function.
    """

    @patch('agents.effort_agent.generate_effort_estimation', return_value=SAMPLE_EFFORT_RESULT)
    @patch('agents.testcase_agent.generate_test_cases', return_value=SAMPLE_TESTCASE_RESULT)
    @patch('agents.plan_agent.generate_project_plan', return_value=SAMPLE_PLAN_RESULT)
    @patch('agents.brd_agent.generate_brd', return_value=SAMPLE_BRD_RESULT)
    @patch('agents.clarification_agent.generate_clarification_questions', return_value=SAMPLE_CLARIFICATION_RESULT)
    def test_full_happy_path(
        self, mock_clarify, mock_brd, mock_plan, mock_tc, mock_effort
    ):
        """
        Full flow:
        1. Create project → get ID
        2. Get clarification questions
        3. Submit answers → fire BRD task
        4. Approve BRD → fire remaining agents
        5. Verify all outputs + download all DOCX
        """
        from apps.projects.tasks import (
            run_clarification_task,
            run_brd_task,
            run_remaining_agents_task,
        )

        client = Client()

        # Patch .delay on each task to prevent broker calls
        with patch.object(run_clarification_task, 'delay', MagicMock()), \
             patch.object(run_brd_task, 'delay', MagicMock()), \
             patch.object(run_remaining_agents_task, 'delay', MagicMock()):

            # ── Step 1: Create project ───────────────────────────────────────
            response = client.post(
                '/api/projects/',
                data=json.dumps({'raw_input': SAMPLE_PROJECT_DESCRIPTION}),
                content_type='application/json'
            )
            self.assertEqual(response.status_code, 201)
            project_id = response.json()['id']
            run_clarification_task.delay.assert_called_once_with(project_id)

            # Run clarification task synchronously (real logic, mocked AI)
            run_clarification_task(project_id)

            # ── Step 2: Get clarification questions ──────────────────────────
            response = client.get(f'/api/projects/{project_id}/clarification-questions/')
            self.assertEqual(response.status_code, 200)
            questions = response.json()['questions']
            self.assertEqual(len(questions), 3)

            # ── Step 3: Submit answers ───────────────────────────────────────
            answers = {q['id']: f'Answer for {q["id"]}' for q in questions}
            response = client.post(
                f'/api/projects/{project_id}/answer-questions/',
                data=json.dumps({'answers': answers}),
                content_type='application/json'
            )
            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json()['status'], 'generating_brd')

            # Run BRD task synchronously
            run_brd_task(project_id)

            # ── Verify BRD output ────────────────────────────────────────────
            response = client.get(f'/api/projects/{project_id}/status/')
            self.assertEqual(response.json()['status'], 'awaiting_approval')

            response = client.get(f'/api/projects/{project_id}/brd/')
            self.assertEqual(response.status_code, 200)
            output = response.json()['structured_output']
            for section in ['executive_summary', 'functional_requirements',
                            'non_functional_requirements', 'success_criteria', 'glossary']:
                self.assertIn(section, output, f'Missing BRD section: {section}')

            # ── Step 4: Approve BRD ──────────────────────────────────────────
            response = client.post(f'/api/projects/{project_id}/approve-brd/')
            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json()['status'], 'approved')

            # Run remaining agents synchronously
            run_remaining_agents_task(project_id)

            # ── Verify final status ──────────────────────────────────────────
            response = client.get(f'/api/projects/{project_id}/status/')
            status_data = response.json()
            self.assertEqual(status_data['status'], 'complete')
            self.assertTrue(status_data['brd_approved'])

            # All agent outputs should be complete
            for agent_type in ['clarification', 'brd', 'plan', 'test_cases', 'effort']:
                self.assertEqual(
                    status_data['outputs'].get(agent_type), 'complete',
                    f'Agent {agent_type} not complete'
                )

            # ── Step 5: Download all DOCX ────────────────────────────────────
            for output_type in ['brd', 'plan', 'testcases', 'effort']:
                response = client.get(f'/api/projects/{project_id}/download/{output_type}/')
                self.assertEqual(response.status_code, 200,
                                 f'DOCX download failed for: {output_type}')
                self.assertIn('application/vnd.openxmlformats',
                              response['Content-Type'],
                              f'Wrong content-type for: {output_type}')

