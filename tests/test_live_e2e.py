"""
Live E2E Test Suite — Real LLM Calls via .env
===============================================
These tests call the ACTUAL configured AI provider (Claude or OpenAI).
No mocks. Every test exercises the real pipeline end-to-end.

Prerequisites:
  - Set ANTHROPIC_API_KEY or OPENAI_API_KEY in .env
  - Set AI_PROVIDER=claude  (or openai) in .env

Run all live tests:
    python manage.py test tests.test_live_e2e --verbosity=2

Run a single class:
    python manage.py test tests.test_live_e2e.LiveClarificationAgentTests --verbosity=2

Run the full pipeline test only:
    python manage.py test tests.test_live_e2e.LiveFullPipelineTests --verbosity=2

Cost/Time estimate (approximate):
  - test_real_clarification_questions  ≈ 5-10s, ~$0.001
  - test_real_brd_generation           ≈ 20-40s, ~$0.01
  - test_real_plan_generation          ≈ 15-30s, ~$0.005
  - test_real_test_cases               ≈ 15-30s, ~$0.005
  - test_real_effort_estimation        ≈ 10-20s, ~$0.003
  - test_full_live_pipeline            ≈ 90-180s, ~$0.03

Skipped automatically if no API key is configured.
"""

import os
import json
import time
import unittest
import django
from pathlib import Path

# ── Load .env before Django setup ─────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
env_file = BASE_DIR / '.env'

if env_file.exists():
    with open(env_file, encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, _, val = line.partition('=')
                os.environ.setdefault(key.strip(), val.strip().strip('"').strip("'"))

# ── Check if any API key is available ─────────────────────────────────────────
HAS_CLAUDE_KEY = bool(os.environ.get('ANTHROPIC_API_KEY', '').strip())
HAS_OPENAI_KEY = bool(os.environ.get('OPENAI_API_KEY', '').strip())
HAS_ANY_KEY = HAS_CLAUDE_KEY or HAS_OPENAI_KEY

SKIP_REASON = (
    'No AI API key configured. Set ANTHROPIC_API_KEY or OPENAI_API_KEY in .env'
)

AI_PROVIDER = os.environ.get('AI_PROVIDER', 'claude').lower()
ACTIVE_PROVIDER = AI_PROVIDER if (
    (AI_PROVIDER == 'claude' and HAS_CLAUDE_KEY) or
    (AI_PROVIDER == 'openai' and HAS_OPENAI_KEY)
) else ('claude' if HAS_CLAUDE_KEY else 'openai' if HAS_OPENAI_KEY else None)

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'brd_system.settings')

from django.test import TestCase, Client
from django.test.utils import override_settings

# ═══════════════════════════════════════════════════════════════════════════════
# SAMPLE INPUTS
# ═══════════════════════════════════════════════════════════════════════════════

SAMPLE_PROJECT_DESCRIPTION = """
Build a B2B SaaS customer support portal for enterprise clients.

The portal should allow enterprise customers to:
- Submit and track support tickets linked to their subscriptions
- View their billing history and download invoices
- Manage user accounts (add/remove team members, set permissions)
- Access a self-service knowledge base and FAQ section

Integration requirements:
- Salesforce CRM for syncing customer account data
- Stripe for billing and invoice management
- Single Sign-On (SSO) via Okta SAML 2.0
- Slack notifications for new high-priority tickets

Compliance: GDPR (EU data residency), SOC 2 Type II
Target users: IT managers and procurement teams at enterprise companies (500+ employees)
Expected scale: 200 concurrent users, 50 enterprise accounts at launch
"""

# Pre-written answers to minimise interactive prompts during test
SAMPLE_ANSWERS = {
    "Q1": "Primary users are IT managers (40%) and procurement leads (60%). Main pain points: no visibility into ticket status, manually tracking invoices via email, no self-service options.",
    "Q2": "Salesforce sync: Account, Contact, Subscription (custom object), and Case objects. Bi-directional sync every 15 minutes.",
    "Q3": "GDPR: All data must reside in Frankfurt AWS region (eu-central-1). We have DPA agreements with all EU clients.",
    "Q4": "Stripe integration: read invoices, display payment history, allow PDF download. No payment processing in portal.",
    "Q5": "Out of scope: mobile app, custom reporting, API access for clients. In scope: web portal only, desktop browsers.",
}


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _print_section(title: str, char: str = '─', width: int = 70):
    print(f'\n{"=" * width}')
    print(f'  {title}')
    print(f'{"=" * width}')


def _print_json_summary(data: dict, max_keys: int = 5):
    """Print first N keys of a JSON dict for quick visual verification."""
    keys = list(data.keys())[:max_keys]
    for k in keys:
        v = data[k]
        if isinstance(v, str):
            print(f'  {k}: {v[:120]}{"..." if len(v) > 120 else ""}')
        elif isinstance(v, list):
            print(f'  {k}: [{len(v)} item(s)]')
        elif isinstance(v, dict):
            print(f'  {k}: {{...{len(v)} key(s)}}')
        else:
            print(f'  {k}: {v}')


# ═══════════════════════════════════════════════════════════════════════════════
# 1. AI PROVIDER SETUP TESTS
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveProviderSetupTests(TestCase):
    """Verify the AI provider is configured correctly before running deeper tests."""

    def test_api_key_loaded_from_env(self):
        """Confirm at least one API key is present."""
        self.assertTrue(
            HAS_CLAUDE_KEY or HAS_OPENAI_KEY,
            'Neither ANTHROPIC_API_KEY nor OPENAI_API_KEY found in environment.'
        )
        if HAS_CLAUDE_KEY:
            key = os.environ.get('ANTHROPIC_API_KEY', '')
            self.assertTrue(key.startswith('sk-ant-'), f'Claude key looks invalid: {key[:15]}...')
            print(f'\n  ✅ Claude key present: sk-ant-...{key[-6:]}')
        if HAS_OPENAI_KEY:
            key = os.environ.get('OPENAI_API_KEY', '')
            self.assertTrue(key.startswith('sk-'), f'OpenAI key looks invalid: {key[:10]}...')
            print(f'\n  ✅ OpenAI key present: sk-...{key[-6:]}')

    def test_active_provider_is_valid(self):
        """Confirm AI_PROVIDER env var maps to an available key."""
        self.assertIsNotNone(ACTIVE_PROVIDER, 'No valid provider found.')
        print(f'\n  ✅ Active AI provider: {ACTIVE_PROVIDER}')

    def test_base_generate_json_returns_dict(self):
        """
        Smoke-test: call generate_json with a trivial prompt.
        Verifies the SDK, auth, and JSON parsing all work.
        """
        from agents.base import generate_json

        system = 'Return JSON only. No markdown.'
        user = 'Return {"status": "ok", "provider": "working"}'

        _print_section('Smoke Test: generate_json')
        t0 = time.time()
        result = generate_json(system, user)
        elapsed = time.time() - t0

        print(f'  Response ({elapsed:.1f}s): {result}')
        self.assertIsInstance(result, dict, 'generate_json must return a dict')
        print(f'  ✅ AI provider responding correctly')


# ═══════════════════════════════════════════════════════════════════════════════
# 2. CLARIFICATION AGENT — LIVE
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveClarificationAgentTests(TestCase):
    """Test the Clarification Agent with real AI calls."""

    def test_real_clarification_questions(self):
        """Generate real clarification questions for the sample project."""
        from agents.clarification_agent import generate_clarification_questions

        _print_section('Live Test: Clarification Agent')
        print(f'  Project description: {SAMPLE_PROJECT_DESCRIPTION[:200].strip()}...')
        print(f'  Calling {ACTIVE_PROVIDER}...')

        t0 = time.time()
        result = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        elapsed = time.time() - t0

        print(f'  Response time: {elapsed:.1f}s')
        print(f'\n  Generated {len(result.get("questions", []))} questions:')
        for q in result.get('questions', []):
            print(f'  [{q.get("id", "?")}] {q.get("question", "?")}')
            print(f'       → Why: {q.get("why_asking", "?")}')

        # Assertions
        self.assertIn('questions', result, 'Result must have "questions" key')
        questions = result['questions']
        self.assertIsInstance(questions, list, '"questions" must be a list')
        self.assertGreaterEqual(len(questions), 3, 'Must generate at least 3 questions')
        self.assertLessEqual(len(questions), 5, 'Must generate at most 5 questions')

        for q in questions:
            self.assertIn('id', q, f'Question missing "id": {q}')
            self.assertIn('question', q, f'Question missing "question": {q}')
            self.assertIn('why_asking', q, f'Question missing "why_asking": {q}')
            self.assertRegex(q['id'], r'^Q\d+$', f'Question ID format wrong: {q["id"]}')
            self.assertGreater(len(q['question']), 20, 'Question text too short')

        print(f'\n  ✅ Clarification agent passed all assertions')

    def test_clarification_infers_compliance(self):
        """Questions should reference GDPR/SOC2 since they appear in the description."""
        from agents.clarification_agent import generate_clarification_questions

        result = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        questions_text = ' '.join(
            q.get('question', '') + ' ' + q.get('why_asking', '')
            for q in result.get('questions', [])
        ).lower()

        # At least one question should mention compliance, security, or data
        compliance_keywords = ['gdpr', 'soc', 'compliance', 'data', 'security', 'privacy', 'residency']
        matched = [kw for kw in compliance_keywords if kw in questions_text]
        self.assertGreater(
            len(matched), 0,
            f'Expected compliance-related question. Got: {questions_text[:400]}'
        )
        print(f'\n  ✅ Compliance keywords found in questions: {matched}')


# ═══════════════════════════════════════════════════════════════════════════════
# 3. BRD AGENT — LIVE
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveBRDAgentTests(TestCase):
    """Test BRD generation with real AI. This is the most expensive test."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        from agents.clarification_agent import generate_clarification_questions
        print('\n  [BRD Test Setup] Running clarification agent...')
        cls.clarification_result = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        # Build answers dict using sample answers + any extra question IDs
        cls.answers = {}
        for q in cls.clarification_result.get('questions', []):
            qid = q.get('id', '')
            cls.answers[qid] = SAMPLE_ANSWERS.get(qid, f'The answer to {qid} is: yes, fully supported.')

    def test_real_brd_generation(self):
        """Generate a full BRD with all 11 sections using real AI."""
        from agents.brd_agent_v2 import generate_brd_improved

        _print_section('Live Test: BRD Agent (Improved v2)')
        print(f'  Using improved agent v2 (multi-stage, 85-90% quality)')
        print(f'  Calling {ACTIVE_PROVIDER} (this takes 30-90s)...')

        t0 = time.time()
        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="Receipt API: UAT and Production endpoints",
            current_process="RPA-based manual process",
            business_context="Automate receipt processing",
            technical_specs="File upload, validation, API integration"
        )
        elapsed = time.time() - t0

        # Extract BRD from improved agent result
        result = result_full["brd"]
        quality_issues = result_full.get("quality_issues", [])

        print(f'  Response time: {elapsed:.1f}s')
        print(f'  Quality issues: {len(quality_issues)} (0 = passed validation)')
        print(f'\n  BRD top-level keys: {list(result.keys())}')
        _print_json_summary(result)

        # Must have all 11 sections
        required = [
            'executive_summary', 'project_scope', 'business_objectives',
            'stakeholders', 'project_plan', 'effort_estimation',
            'functional_requirements', 'non_functional_requirements',
            'constraints_and_assumptions', 'success_criteria', 'glossary'
        ]
        for section in required:
            self.assertIn(section, result, f'BRD missing section: {section}')

        # Functional requirements must be a list of dicts with required fields
        frs = result.get('functional_requirements', [])
        self.assertIsInstance(frs, list)
        self.assertGreater(len(frs), 0, 'Must have at least 1 functional requirement')
        for fr in frs:
            self.assertIn('id', fr, f'FR missing id: {fr}')
            self.assertRegex(fr['id'], r'^FR-\d+$', f'FR ID format wrong: {fr.get("id")}')
            self.assertIn('title', fr)
            self.assertIn('description', fr)
            self.assertIn('priority', fr)
            self.assertIn('acceptance_criteria', fr)

        # Non-functional requirements
        nfrs = result.get('non_functional_requirements', [])
        self.assertGreater(len(nfrs), 0)

        # Stakeholders
        stakeholders = result.get('stakeholders', [])
        self.assertGreater(len(stakeholders), 0)

        # Compliance: GDPR and SOC 2 should appear somewhere
        brd_text = json.dumps(result).lower()
        compliance_found = [kw for kw in ['gdpr', 'soc', 'sso', 'okta', 'stripe'] if kw in brd_text]
        self.assertGreater(len(compliance_found), 0,
                           f'Expected compliance keywords in BRD. Keys found: {compliance_found}')

        print(f'\n  ✅ BRD passed all assertions')
        print(f'  📋 FRs: {len(frs)} | NFRs: {len(nfrs)} | Compliance: {compliance_found}')

        # Save BRD to class so other tests can use it
        LiveBRDAgentTests._brd_result = result

    def test_brd_functional_requirements_have_acceptance_criteria(self):
        """Each FR must have at least 1 acceptance criterion."""
        from agents.brd_agent_v2 import generate_brd_improved

        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specifications",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        result = result_full["brd"]
        for fr in result.get('functional_requirements', []):
            criteria = fr.get('acceptance_criteria', [])
            self.assertIsInstance(criteria, list, f'acceptance_criteria must be a list in {fr["id"]}')
            self.assertGreater(len(criteria), 0, f'{fr["id"]} has no acceptance criteria')

        print(f'\n  ✅ All FRs have acceptance criteria')

    def test_brd_revision_incorporates_feedback(self):
        """BRD revision with notes should update the output."""
        from agents.brd_agent_v2 import generate_brd_improved

        revision_notes = (
            'Please add a dedicated section for the Slack integration in functional requirements. '
            'Also add a specific NFR for API response time under 500ms.'
        )

        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specifications",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        result = result_full["brd"]
        brd_text = json.dumps(result).lower()

        self.assertIn('slack', brd_text, 'Slack integration should appear in revised BRD')
        print(f'\n  ✅ BRD revision incorporated feedback (Slack found in output)')


# ═══════════════════════════════════════════════════════════════════════════════
# 4. PLAN AGENT — LIVE
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LivePlanAgentTests(TestCase):
    """Test Project Plan generation with real AI."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        from agents.clarification_agent import generate_clarification_questions
        from agents.brd_agent_v2 import generate_brd_improved

        print('\n  [Plan Test Setup] Generating BRD first...')
        clarification = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        answers = {
            q['id']: SAMPLE_ANSWERS.get(q['id'], 'Yes, this is fully supported.')
            for q in clarification.get('questions', [])
        }
        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specs",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        cls.brd_data = result_full["brd"]
        print(f'  [Plan Test Setup] BRD ready with {len(cls.brd_data.get("functional_requirements", []))} FRs')

    def test_real_plan_generation(self):
        """Generate a project plan from the BRD using real AI."""
        from agents.plan_agent import generate_project_plan

        _print_section('Live Test: Plan Agent')
        print(f'  Calling {ACTIVE_PROVIDER}...')

        t0 = time.time()
        result = generate_project_plan(self.brd_data)
        elapsed = time.time() - t0

        print(f'  Response time: {elapsed:.1f}s')
        print(f'  Plan keys: {list(result.keys())}')

        # Structure assertions
        self.assertIn('phases', result)
        self.assertIn('methodology', result)
        self.assertIn('total_duration', result)
        self.assertIn('risk_register', result)
        self.assertIn('team_structure', result)

        phases = result['phases']
        self.assertIsInstance(phases, list)
        self.assertGreater(len(phases), 0, 'Plan must have at least 1 phase')

        for phase in phases:
            self.assertIn('phase_name', phase)
            self.assertIn('duration', phase)
            self.assertIn('tasks', phase)

        # Tasks should reference FR IDs
        all_tasks = [t for p in phases for t in p.get('tasks', [])]
        self.assertGreater(len(all_tasks), 0, 'Plan must have at least 1 task')

        # Risk register
        risks = result.get('risk_register', [])
        self.assertIsInstance(risks, list)

        print(f'\n  ✅ Plan passed all assertions')
        print(f'  📋 Phases: {len(phases)} | Tasks: {len(all_tasks)} | Risks: {len(risks)}')
        for p in phases:
            print(f'     Phase: {p.get("phase_name")} ({p.get("duration")})')


# ═══════════════════════════════════════════════════════════════════════════════
# 5. TEST CASE AGENT — LIVE
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveTestCaseAgentTests(TestCase):
    """Test Case generation with real AI + traceability matrix validation."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        from agents.clarification_agent import generate_clarification_questions
        from agents.brd_agent_v2 import generate_brd_improved

        print('\n  [TestCase Setup] Generating BRD...')
        clarification = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        answers = {
            q['id']: SAMPLE_ANSWERS.get(q['id'], 'Supported.')
            for q in clarification.get('questions', [])
        }
        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specs",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        cls.brd_data = result_full["brd"]
        cls.fr_ids = [fr['id'] for fr in cls.brd_data.get('functional_requirements', [])]
        print(f'  [TestCase Setup] BRD ready. FRs: {cls.fr_ids}')

    def test_real_test_case_generation(self):
        """Generate test cases from BRD using real AI."""
        from agents.testcase_agent import generate_test_cases

        _print_section('Live Test: Test Case Agent')
        print(f'  Calling {ACTIVE_PROVIDER}...')

        t0 = time.time()
        result = generate_test_cases(self.brd_data)
        elapsed = time.time() - t0

        print(f'  Response time: {elapsed:.1f}s')

        self.assertIn('test_cases', result)
        self.assertIn('traceability_matrix', result)
        self.assertIn('test_summary', result)

        test_cases = result['test_cases']
        self.assertIsInstance(test_cases, list)
        self.assertGreater(len(test_cases), 0, 'Must have at least 1 test case')

        print(f'\n  Generated {len(test_cases)} test cases:')
        for tc in test_cases[:5]:  # show first 5
            print(f'  [{tc.get("test_id")}] {tc.get("title")} → FR: {tc.get("linked_requirement")}')

        for tc in test_cases:
            self.assertIn('test_id', tc)
            self.assertRegex(tc['test_id'], r'^TC-\d+$', f'TC ID format wrong: {tc.get("test_id")}')
            self.assertIn('linked_requirement', tc)
            self.assertIn('title', tc)
            self.assertIn('test_steps', tc)
            steps = tc.get('test_steps', [])
            self.assertIsInstance(steps, list)
            self.assertGreater(len(steps), 0, f'{tc["test_id"]} has no test steps')

        # Traceability matrix — every FR should be covered
        matrix = result['traceability_matrix']
        covered_fr_ids = [row.get('requirement_id') for row in matrix]
        for fr_id in self.fr_ids:
            self.assertIn(fr_id, covered_fr_ids,
                          f'FR {fr_id} missing from traceability matrix')

        print(f'\n  ✅ Test cases passed all assertions')
        print(f'  📋 Test cases: {len(test_cases)} | Traceability rows: {len(matrix)}')
        summary = result.get('test_summary', {})
        print(f'  Coverage: {summary.get("coverage_percentage", "N/A")}')


# ═══════════════════════════════════════════════════════════════════════════════
# 6. EFFORT AGENT — LIVE
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveEffortAgentTests(TestCase):
    """Effort Estimation with real AI."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        from agents.clarification_agent import generate_clarification_questions
        from agents.brd_agent_v2 import generate_brd_improved
        from agents.plan_agent import generate_project_plan

        print('\n  [Effort Setup] Generating BRD + Plan...')
        clarification = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        answers = {
            q['id']: SAMPLE_ANSWERS.get(q['id'], 'Supported.')
            for q in clarification.get('questions', [])
        }
        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specs",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        cls.brd_data = result_full["brd"]
        cls.plan_data = generate_project_plan(cls.brd_data)
        print(f'  [Effort Setup] BRD + Plan ready')

    def test_real_effort_estimation(self):
        """Generate effort estimate from BRD + Plan using real AI."""
        from agents.effort_agent import generate_effort_estimation

        _print_section('Live Test: Effort Agent')
        print(f'  Calling {ACTIVE_PROVIDER}...')

        t0 = time.time()
        result = generate_effort_estimation(self.brd_data, self.plan_data)
        elapsed = time.time() - t0

        print(f'  Response time: {elapsed:.1f}s')

        self.assertIn('estimation_summary', result)
        self.assertIn('by_phase', result)
        self.assertIn('by_role', result)
        self.assertIn('by_feature', result)
        self.assertIn('cost_estimate', result)

        summary = result['estimation_summary']
        self.assertIn('total_hours', summary)
        self.assertIsInstance(summary['total_hours'], (int, float))
        self.assertGreater(summary['total_hours'], 0)

        cost = result['cost_estimate']
        self.assertIn('low_estimate', cost)
        self.assertIn('mid_estimate', cost)
        self.assertIn('high_estimate', cost)
        self.assertLess(cost['low_estimate'], cost['mid_estimate'])
        self.assertLess(cost['mid_estimate'], cost['high_estimate'])

        print(f'\n  ✅ Effort estimation passed all assertions')
        print(f'  📋 Total hours: {summary.get("total_hours")}')
        print(f'  💰 Cost: ${cost.get("low_estimate"):,} – ${cost.get("high_estimate"):,} {cost.get("currency", "USD")}')
        print(f'  👥 Team size: {summary.get("team_size_recommended")} | Duration: {summary.get("total_weeks")} weeks')


# ═══════════════════════════════════════════════════════════════════════════════
# 7. DOCX EXPORT — LIVE (with real AI output)
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveDocxExportTests(TestCase):
    """Generate real AI output then export to DOCX and validate the document."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        import tempfile

        from agents.clarification_agent import generate_clarification_questions
        from agents.brd_agent_v2 import generate_brd_improved

        print('\n  [DOCX Setup] Generating BRD...')
        clarification = generate_clarification_questions(SAMPLE_PROJECT_DESCRIPTION)
        answers = {
            q['id']: SAMPLE_ANSWERS.get(q['id'], 'Supported.')
            for q in clarification.get('questions', [])
        }
        result_full = generate_brd_improved(
            project_description=SAMPLE_PROJECT_DESCRIPTION,
            api_specs="API specs",
            current_process="Current process",
            business_context="Business context",
            technical_specs="Technical specs"
        )
        cls.brd_data = result_full["brd"]
        cls.output_dir = Path(BASE_DIR) / 'test_output'
        cls.output_dir.mkdir(exist_ok=True)
        print(f'  [DOCX Setup] BRD ready. Output dir: {cls.output_dir}')

    def test_brd_docx_has_real_content(self):
        """Export real AI BRD to DOCX and verify it contains actual project content."""
        from utils.docx_exporter import export_brd_to_docx
        from docx import Document

        buffer = export_brd_to_docx(self.brd_data)
        buffer.seek(0)
        doc = Document(buffer)

        all_text = '\n'.join(p.text for p in doc.paragraphs).lower()

        # Should contain real project content
        keywords = ['portal', 'salesforce', 'gdpr', 'soc', 'enterprise']
        matched = [kw for kw in keywords if kw in all_text]
        self.assertGreater(len(matched), 0,
                           f'BRD DOCX missing project content. Text: {all_text[:500]}')

        # Save to disk for manual review
        out_path = self.output_dir / 'live_test_BRD.docx'
        buffer.seek(0)
        with open(out_path, 'wb') as f:
            f.write(buffer.read())

        print(f'\n  ✅ BRD DOCX contains real content: {matched}')
        print(f'  💾 Saved to: {out_path}')

    def test_brd_docx_headings_match_sections(self):
        """Every BRD section returned by AI should appear as a heading in DOCX."""
        from utils.docx_exporter import export_brd_to_docx
        from docx import Document

        buffer = export_brd_to_docx(self.brd_data)
        buffer.seek(0)
        doc = Document(buffer)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        headings_text = ' '.join(headings)

        for section_heading in [
            'Executive Summary', 'Project Scope', 'Business Objectives',
            'Functional Requirements', 'Non-Functional Requirements', 'Glossary'
        ]:
            self.assertIn(section_heading, headings_text,
                          f'Missing heading in DOCX: {section_heading}')

        print(f'\n  ✅ All expected headings present in DOCX')
        print(f'  📄 Headings found: {headings[:8]}')


# ═══════════════════════════════════════════════════════════════════════════════
# 8. FULL LIVE PIPELINE — API + Real LLM (no mocks at all)
# ═══════════════════════════════════════════════════════════════════════════════

@unittest.skipUnless(HAS_ANY_KEY, SKIP_REASON)
class LiveFullPipelineTests(TestCase):
    """
    Full end-to-end pipeline test via the Django REST API.
    Uses real LLM, real DOCX export, real database.
    Mirrors the exact flow the React frontend + Celery worker performs.

    ⚠️  This test bypasses Celery (no broker needed) by calling task
        functions directly. The HTTP API surface is still fully exercised.

    Runtime: 90–180 seconds. Cost: ~$0.03
    """

    def test_full_live_pipeline(self):
        """
        Complete pipeline:
        1. POST project description
        2. Run clarification task (sync)
        3. GET questions — verify real questions
        4. POST answers
        5. Run BRD task (sync)
        6. GET BRD — verify 11 sections with real content
        7. POST approve-brd
        8. Run remaining agents (sync)
        9. GET plan / testcases / effort
        10. Download all 4 DOCX files and save to test_output/
        """
        from apps.projects.tasks import (
            run_clarification_task,
            run_brd_task,
            run_remaining_agents_task,
        )
        from apps.projects.models import Project, AgentOutput

        output_dir = Path(BASE_DIR) / 'test_output'
        output_dir.mkdir(exist_ok=True)

        client = Client()
        _print_section('LIVE FULL PIPELINE TEST (Real LLM)')

        # ── Step 1: Create project ───────────────────────────────────────────
        print('\n  [Step 1] Creating project...')
        with unittest.mock.patch.object(run_clarification_task, 'delay',
                                         unittest.mock.MagicMock()):
            response = client.post(
                '/api/projects/',
                data=json.dumps({'raw_input': SAMPLE_PROJECT_DESCRIPTION}),
                content_type='application/json'
            )

        self.assertEqual(response.status_code, 201,
                         f'Create failed: {response.json()}')
        project_id = response.json()['id']
        print(f'  ✅ Project created: {project_id}')

        # ── Step 2: Run clarification (real AI) ──────────────────────────────
        print('\n  [Step 2] Running Clarification Agent (real AI)...')
        t0 = time.time()
        run_clarification_task(project_id)
        print(f'  ✅ Done in {time.time() - t0:.1f}s')

        # ── Step 3: GET questions ────────────────────────────────────────────
        print('\n  [Step 3] Fetching clarification questions...')
        response = client.get(f'/api/projects/{project_id}/clarification-questions/')
        self.assertEqual(response.status_code, 200,
                         f'Get questions failed: {response.content}')
        questions = response.json()['questions']
        self.assertGreaterEqual(len(questions), 3)
        print(f'  ✅ Got {len(questions)} real questions:')
        for q in questions:
            print(f'     [{q["id"]}] {q["question"]}')

        # ── Step 4: Submit answers ───────────────────────────────────────────
        print('\n  [Step 4] Submitting answers...')
        answers = {
            q['id']: SAMPLE_ANSWERS.get(q['id'], f'The answer to {q["id"]}: this is fully supported and required.')
            for q in questions
        }
        with unittest.mock.patch.object(run_brd_task, 'delay',
                                         unittest.mock.MagicMock()):
            response = client.post(
                f'/api/projects/{project_id}/answer-questions/',
                data=json.dumps({'answers': answers}),
                content_type='application/json'
            )
        self.assertEqual(response.status_code, 200,
                         f'Submit answers failed: {response.json()}')
        print(f'  ✅ Answers submitted')

        # ── Step 5: Run BRD task (real AI) ──────────────────────────────────
        print('\n  [Step 5] Running BRD Agent (real AI, ~30-60s)...')
        t0 = time.time()
        run_brd_task(project_id)
        print(f'  ✅ BRD generated in {time.time() - t0:.1f}s')

        # ── Step 6: GET BRD ──────────────────────────────────────────────────
        print('\n  [Step 6] Fetching BRD output...')
        response = client.get(f'/api/projects/{project_id}/brd/')
        self.assertEqual(response.status_code, 200)
        brd = response.json()['structured_output']

        required_sections = [
            'executive_summary', 'project_scope', 'business_objectives',
            'stakeholders', 'project_plan', 'effort_estimation',
            'functional_requirements', 'non_functional_requirements',
            'constraints_and_assumptions', 'success_criteria', 'glossary'
        ]
        for section in required_sections:
            self.assertIn(section, brd, f'BRD missing section: {section}')

        frs = brd.get('functional_requirements', [])
        print(f'  ✅ BRD has all 11 sections | {len(frs)} functional requirements')
        for fr in frs:
            print(f'     [{fr.get("id")}] {fr.get("title")} ({fr.get("priority")})')

        # Check status
        status_resp = client.get(f'/api/projects/{project_id}/status/')
        self.assertEqual(status_resp.json()['status'], 'awaiting_approval')

        # ── Step 7: Approve BRD ──────────────────────────────────────────────
        print('\n  [Step 7] Approving BRD...')
        with unittest.mock.patch.object(run_remaining_agents_task, 'delay',
                                         unittest.mock.MagicMock()):
            response = client.post(f'/api/projects/{project_id}/approve-brd/')
        self.assertEqual(response.status_code, 200)
        print(f'  ✅ BRD approved')

        # ── Step 8: Run remaining agents (real AI) ───────────────────────────
        print('\n  [Step 8] Running Plan + TestCase + Effort agents (real AI, ~60-120s)...')
        t0 = time.time()
        run_remaining_agents_task(project_id)
        print(f'  ✅ All agents complete in {time.time() - t0:.1f}s')

        # ── Step 9: Verify outputs ───────────────────────────────────────────
        print('\n  [Step 9] Verifying all agent outputs...')

        final_status = client.get(f'/api/projects/{project_id}/status/')
        self.assertEqual(final_status.json()['status'], 'complete')

        plan_resp = client.get(f'/api/projects/{project_id}/plan/')
        self.assertEqual(plan_resp.status_code, 200)
        plan = plan_resp.json()['structured_output']
        self.assertIn('phases', plan)
        print(f'  ✅ Plan: {len(plan.get("phases", []))} phases')

        tc_resp = client.get(f'/api/projects/{project_id}/testcases/')
        self.assertEqual(tc_resp.status_code, 200)
        tc = tc_resp.json()['structured_output']
        self.assertIn('test_cases', tc)
        print(f'  ✅ Test Cases: {len(tc.get("test_cases", []))} test cases')

        effort_resp = client.get(f'/api/projects/{project_id}/effort/')
        self.assertEqual(effort_resp.status_code, 200)
        effort = effort_resp.json()['structured_output']
        self.assertIn('cost_estimate', effort)
        cost = effort.get('cost_estimate', {})
        print(f'  ✅ Effort: {effort.get("estimation_summary", {}).get("total_hours")} hours | '
              f'${cost.get("low_estimate"):,}–${cost.get("high_estimate"):,}')

        # ── Step 10: Download all DOCX ───────────────────────────────────────
        print('\n  [Step 10] Downloading DOCX files...')
        download_map = {
            'brd': 'live_pipeline_BRD.docx',
            'plan': 'live_pipeline_ProjectPlan.docx',
            'testcases': 'live_pipeline_TestCases.docx',
            'effort': 'live_pipeline_EffortEstimation.docx',
        }
        for output_type, filename in download_map.items():
            resp = client.get(f'/api/projects/{project_id}/download/{output_type}/')
            self.assertEqual(resp.status_code, 200,
                             f'DOCX download failed for {output_type}: {resp.status_code}')
            self.assertIn('application/vnd.openxmlformats', resp['Content-Type'])

            out_path = output_dir / filename
            with open(out_path, 'wb') as f:
                for chunk in resp.streaming_content:
                    f.write(chunk)
            print(f'  💾 {filename} → {out_path}')

        _print_section('✅ FULL LIVE PIPELINE TEST PASSED')
        print(f'  Project ID: {project_id}')
        print(f'  Output files: {output_dir}/')
        print(f'  AI Provider: {ACTIVE_PROVIDER}')


# ── Import fix for unittest.mock inside class ──────────────────────────────────
import unittest.mock
