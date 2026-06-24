"""
ABHFL-Branded DOCX Export Utility.

Generates Word documents that exactly match the Aditya Birla Housing Finance Ltd (ABHFL)
BRD format — same logo, colors, fonts, and structure as the reference BRDs in the BRDs/ folder.

Brand tokens:
  - Primary Green  : #009900  (Title, Heading 1)
  - Primary Blue   : #2F5496  (Heading 2, subtitle)
  - Body Blue      : #2E74B5  (accent text)
  - Body Text      : #374151  (Segoe UI)
  - Table Header BG: #F2F2F2
  - Table Header FG: #009900 bold
"""

from io import BytesIO
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ─── Brand Constants ─────────────────────────────────────────────────────────

ABHFL_GREEN  = RGBColor(0x00, 0x99, 0x00)   # #009900
ABHFL_RED    = RGBColor(0xC0, 0x00, 0x00)   # #C00000 (ABHFL brand red from logo)
ABHFL_BLUE   = RGBColor(0x2F, 0x54, 0x96)   # #2F5496
ABHFL_BLUE2  = RGBColor(0x2E, 0x74, 0xB5)   # #2E74B5
ABHFL_BODY   = RGBColor(0x37, 0x41, 0x51)   # #374151
ABHFL_GREY   = "F2F2F2"                      # table header background (hex string)
ABHFL_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

LOGO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "media", "brd_assets", "ey_image_14.png"
)

FONT_HEADING = "Calibri"
FONT_BODY    = "Segoe UI"


# ─── Low-level XML Helpers ────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str, color_hex: str = "auto"):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), color_hex)
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(table):
    """Apply thin grey borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _add_toc(doc: Document):
    """Insert a Word TOC field that updates on first open."""
    para = doc.add_paragraph()
    run = para.add_run()
    for tag, ftype in [
        ("w:fldChar", "begin"),
        ("w:instrText", None),
        ("w:fldChar", "separate"),
        ("w:fldChar", "end"),
    ]:
        el = OxmlElement(tag)
        if tag == "w:instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = 'TOC \\o "1-3" \\h \\z \\u'
        elif tag == "w:fldChar":
            el.set(qn("w:fldCharType"), ftype)
        run._r.append(el)


def _paragraph_spacing(para, before_pt: int = 0, after_pt: int = 6):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before_pt * 20))
    spacing.set(qn("w:after"),  str(after_pt * 20))
    pPr.append(spacing)


# ─── Shared Layout Helpers ────────────────────────────────────────────────────

def _setup_page(doc: Document):
    """Apply ABHFL page margins (1 inch all sides, A4)."""
    section = doc.sections[0]
    section.page_width  = Emu(7560310)   # A4 width
    section.page_height = Emu(10692130)  # A4 height
    margin = Inches(1)
    section.top_margin    = margin
    section.bottom_margin = margin
    section.left_margin   = margin
    section.right_margin  = margin


def _add_logo_header(doc: Document):
    """No-op: logo is only on the cover page, not in the running header."""
    pass  # intentionally empty — header kept clean on all pages


def _add_footer(doc: Document, project_name: str = ""):
    """Add page numbers and company name to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    left_run = fp.add_run("Aditya Birla Housing Finance Ltd.   |   ")
    left_run.font.name = FONT_BODY
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = ABHFL_BODY

    # Page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    page_run = fp.add_run()
    page_run.font.name = FONT_BODY
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = ABHFL_BODY
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)
    page_run._r.append(fldChar3)


def _add_cover_page(doc: Document, title: str, project_name: str, version: str = "1.0", date: str = ""):
    """
    Add ABHFL-branded cover page matching the EY BRD layout:
      - Logo banner at top (full width)
      - Centered: document type title, project name, date, version
    No logo appears on any other page.
    """
    # Logo — full width, top of cover only
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_para.add_run()
        logo_run.add_picture(LOGO_PATH, width=Inches(6.0))
        _paragraph_spacing(logo_para, before_pt=0, after_pt=60)
    else:
        name_para = doc.add_paragraph()
        r = name_para.add_run("Aditya Birla Housing Finance Ltd.")
        r.bold = True
        r.font.color.rgb = ABHFL_RED
        r.font.name = FONT_HEADING
        r.font.size = Pt(16)
        _paragraph_spacing(name_para, before_pt=0, after_pt=60)

    # Centered document type title (e.g. "Business Requirements Document (BRD)")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = FONT_BODY
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = ABHFL_BLUE
    _paragraph_spacing(title_para, before_pt=0, after_pt=6)

    # Project name — centered, bold blue
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(project_name)
    sub_run.bold = True
    sub_run.font.name = FONT_BODY
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = ABHFL_BLUE
    _paragraph_spacing(sub_para, before_pt=0, after_pt=6)

    # Date — centered
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date or "")
    date_run.font.name = FONT_BODY
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = ABHFL_BODY
    _paragraph_spacing(date_para, before_pt=0, after_pt=4)

    # Version — centered
    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_run = ver_para.add_run(f"Version {version}")
    ver_run.font.name = FONT_BODY
    ver_run.font.size = Pt(11)
    ver_run.font.color.rgb = ABHFL_BODY
    _paragraph_spacing(ver_para, before_pt=0, after_pt=0)

    doc.add_page_break()


def _add_h1(doc: Document, text: str):
    """ABHFL Heading 1 — Calibri Bold #009900."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(14)
    run.font.color.rgb = ABHFL_GREEN
    _paragraph_spacing(para, before_pt=14, after_pt=4)
    return para


def _add_h2(doc: Document, text: str):
    """ABHFL Heading 2 — Calibri Bold #2F5496."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(12)
    run.font.color.rgb = ABHFL_BLUE
    _paragraph_spacing(para, before_pt=10, after_pt=2)
    return para


def _add_body(doc: Document, text: str):
    """Normal body text — Segoe UI #374151."""
    para = doc.add_paragraph()
    run = para.add_run(str(text))
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = ABHFL_BODY
    _paragraph_spacing(para, before_pt=0, after_pt=4)
    return para


def _add_label_value(doc: Document, label: str, value: str):
    """Bold green label + body text value on same line."""
    para = doc.add_paragraph()
    r1 = para.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = FONT_BODY
    r1.font.size = Pt(10)
    r1.font.color.rgb = ABHFL_GREEN
    r2 = para.add_run(str(value))
    r2.font.name = FONT_BODY
    r2.font.size = Pt(10)
    r2.font.color.rgb = ABHFL_BODY
    _paragraph_spacing(para, before_pt=0, after_pt=3)
    return para


def _add_bullet(doc: Document, text: str):
    """ABHFL styled bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(str(text))
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = ABHFL_BODY
    _paragraph_spacing(para, before_pt=0, after_pt=2)
    return para


def _add_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    """
    ABHFL-branded table:
      - Header row: #F2F2F2 background, #009900 bold text
      - Body rows: alternating white / #F9F9F9
      - Thin grey borders
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_cell_borders(table)

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_shading(cell, ABHFL_GREY)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.name = FONT_HEADING
        run.font.size = Pt(9)
        run.font.color.rgb = ABHFL_GREEN

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = "FFFFFF" if ri % 2 == 0 else "F9F9F9"
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            _set_cell_shading(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val) if val is not None else "")
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.color.rgb = ABHFL_BODY

    doc.add_paragraph()  # breathing space after table
    return table


# ─── BRD Exporter ────────────────────────────────────────────────────────────

def export_brd_to_docx(structured_output: dict, toc_sections: list[dict] = None, project_name: str = None, version: str = None) -> BytesIO:
    """
    Export BRD JSON to a fully ABHFL-branded Word document.
    """
    from datetime import date as dt

    doc = Document()
    _setup_page(doc)
    _add_logo_header(doc)
    _add_footer(doc)

    p_name = project_name or structured_output.get("project_name") or ""
    # If the project name is identical to the document title, don't repeat it
    if p_name.lower() == "business requirements document":
        p_name = ""
        
    ver = version or structured_output.get("version", "1.0")
    today = dt.today().strftime("%B %Y")

    _add_cover_page(
        doc,
        title="Business Requirements Document",
        project_name=p_name,
        version=ver,
        date=structured_output.get("document_date", today),
    )

    # ── Table of Contents ──
    _add_h1(doc, "Table of Contents")
    
    if not toc_sections:
        keys = ["executive_summary", "project_scope", "business_objectives", "stakeholders", 
                "project_plan", "effort_estimation", "functional_requirements", 
                "non_functional_requirements", "constraints_and_assumptions", 
                "success_criteria", "glossary"]
        toc_sections = [{"key": k, "label": k.replace('_', ' ').title()} for k in keys]
        
    for idx, sec in enumerate(toc_sections, start=1):
        key = sec['key']
        label = sec.get('label', key.replace('_', ' ').title())
        _add_body(doc, f"{idx}. {label}")
        
    doc.add_page_break()

    idx = 0
    for idx, sec in enumerate(toc_sections, start=1):
        key = sec['key']
        label = sec.get('label', key.replace('_', ' ').title())
        content = structured_output.get(key)
        
        # Add heading
        _add_h1(doc, f"{idx}. {label}")
        
        if not content:
            _add_body(doc, "Content not provided.")
            continue
            
        if key == "executive_summary":
            _add_body(doc, content)
        elif key == "project_scope":
            _add_h2(doc, f"{idx}.1 In Scope")
            for item in content.get("in_scope", []):
                _add_bullet(doc, item)
            _add_h2(doc, f"{idx}.2 Out of Scope")
            for item in content.get("out_of_scope", []):
                _add_bullet(doc, item)
        elif key == "business_objectives":
            for obj in content:
                _add_h2(doc, f"{obj.get('id', '')} — {obj.get('objective', '')}")
                _add_label_value(doc, "Metric", obj.get("metric", ""))
                _add_label_value(doc, "Target", obj.get("target", ""))
        elif key == "stakeholders":
            _add_table(
                doc,
                headers=["Role", "Responsibilities", "Interest Level", "Influence Level"],
                rows=[
                    [s.get("role", ""), s.get("responsibilities", ""),
                     s.get("interest_level", ""), s.get("influence_level", "")]
                    for s in content
                ],
                col_widths=[1.5, 3.0, 1.0, 1.0],
            )
        elif key == "project_plan":
            for phase in content.get("phases", []):
                _add_h2(doc, f"{phase.get('phase', '')}  ({phase.get('duration', '')})")
                _add_label_value(doc, "Deliverables", "")
                for d in phase.get("deliverables", []):
                    _add_bullet(doc, d)
                _add_label_value(doc, "Milestones", "")
                for m in phase.get("milestones", []):
                    _add_bullet(doc, m)
        elif key == "effort_estimation":
            _add_label_value(doc, "Total Estimated Hours", str(content.get("total_estimated_hours", "TBD")))
            _add_body(doc, content.get("summary", ""))
            breakdown = content.get("breakdown", [])
            if breakdown:
                _add_table(
                    doc,
                    headers=["Component", "Hours", "Complexity"],
                    rows=[[b.get("component", ""), b.get("hours", ""), b.get("complexity", "")] for b in breakdown],
                    col_widths=[3.0, 1.5, 2.0],
                )
        elif key == "functional_requirements":
            for req in content:
                _add_h2(doc, f"{req.get('id', '')} — {req.get('title', '')}")
                _add_label_value(doc, "Priority", req.get("priority", ""))
                _add_body(doc, req.get("description", ""))
                _add_label_value(doc, "Acceptance Criteria", "")
                for ac in req.get("acceptance_criteria", []):
                    _add_bullet(doc, ac)
                compliance = req.get("compliance_notes", "")
                if compliance and compliance != "N/A":
                    _add_label_value(doc, "Compliance", compliance)
        elif key == "non_functional_requirements":
            _add_table(
                doc,
                headers=["ID", "Category", "Requirement", "Metric", "Priority"],
                rows=[
                    [n.get("id", ""), n.get("category", ""), n.get("requirement", ""),
                     n.get("metric", ""), n.get("priority", "")]
                    for n in content
                ],
                col_widths=[0.7, 1.2, 2.5, 1.5, 0.7],
            )
        elif key in ["constraints_and_assumptions", "assumptions_and_dependencies"]:
            _add_h2(doc, f"{idx}.1 Constraints")
            for con in content.get("constraints", []):
                _add_h2(doc, f"{con.get('id', '')} — {con.get('description', '')}")
                _add_label_value(doc, "Impact", con.get("impact", ""))
            _add_h2(doc, f"{idx}.2 Assumptions")
            for ass in content.get("assumptions", []):
                _add_h2(doc, f"{ass.get('id', '')} — {ass.get('description', '')}")
                _add_label_value(doc, "Risk if Wrong", ass.get("risk_if_wrong", ""))
        elif key == "success_criteria":
            for sc in content:
                _add_h2(doc, f"{sc.get('id', '')} — {sc.get('criterion', '')}")
                _add_label_value(doc, "Measurement", sc.get("measurement_method", ""))
                _add_label_value(doc, "Target", sc.get("target", ""))
        elif key == "glossary":
            _add_table(
                doc,
                headers=["Term", "Definition"],
                rows=[[g.get("term", ""), g.get("definition", "")] for g in content],
                col_widths=[2.0, 4.5],
            )
        elif key == "integration_requirements":
            for sc in content:
                _add_h2(doc, f"{sc.get('system', '')} — {sc.get('integration_type', '')}")
                _add_body(doc, sc.get("description", ""))
        elif key == "risks_and_mitigations":
            for sc in content:
                _add_h2(doc, f"{sc.get('id', '')} — {sc.get('description', '')}")
                _add_label_value(doc, "Probability", sc.get("probability", ""))
                _add_label_value(doc, "Impact", sc.get("impact", ""))
                _add_label_value(doc, "Mitigation", sc.get("mitigation", ""))
        elif key == "sign_off_matrix":
            _add_table(
                doc,
                headers=["Role", "Name", "Sign-off Date"],
                rows=[[g.get("role", ""), g.get("name", ""), g.get("sign_off_date", "")] for g in content],
                col_widths=[2.0, 2.5, 2.0],
            )
        else:
            # Custom section fallback
            if isinstance(content, dict) and "content_blocks" in content:
                for block in content.get("content_blocks", []):
                    btype = block.get("type")
                    if btype == "paragraph":
                        _add_body(doc, block.get("text", ""))
                    elif btype == "list":
                        for item in block.get("items", []):
                            _add_bullet(doc, str(item))
                    elif btype == "table":
                        headers = block.get("headers", [])
                        rows = block.get("rows", [])
                        if headers and rows:
                            # Try to make col widths proportional
                            num_cols = len(headers)
                            col_widths = [6.0 / num_cols] * num_cols if num_cols else []
                            _add_table(doc, headers=headers, rows=rows, col_widths=col_widths)
            elif isinstance(content, str):
                _add_body(doc, content)
            elif isinstance(content, list):
                for item in content:
                    _add_bullet(doc, str(item))
            elif isinstance(content, dict):
                for k, v in content.items():
                    if k != "title":
                        _add_h2(doc, k.replace('_', ' ').title())
                    if isinstance(v, list):
                        for item in v:
                            _add_bullet(doc, str(item))
                    elif isinstance(v, str):
                        _add_body(doc, str(v))
            else:
                _add_body(doc, str(content))

    # ── Document Control ──
    doc_idx = idx + 1
    _add_h1(doc, f"{doc_idx}. Document Control")
    _add_table(
        doc,
        headers=["Date", "Version", "Description", "Author"],
        rows=[[today, structured_output.get("version", "1.0"), "Initial AI-generated BRD", "AI Agent"]],
        col_widths=[1.2, 0.8, 3.5, 1.0],
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Project Plan Exporter ────────────────────────────────────────────────────

def export_plan_to_docx(structured_output: dict) -> BytesIO:
    """Export Project Plan JSON to a fully ABHFL-branded Word document."""
    from datetime import date as dt

    doc = Document()
    _setup_page(doc)
    _add_logo_header(doc)
    _add_footer(doc)

    project_name = structured_output.get("project_name", "Project")
    today = dt.today().strftime("%B %Y")
    _add_cover_page(doc, "Project Plan", project_name, date=today)

    _add_h1(doc, "Project Summary")
    _add_body(doc, structured_output.get("project_summary", ""))
    _add_label_value(doc, "Methodology", structured_output.get("methodology", ""))
    _add_label_value(doc, "Total Duration", structured_output.get("total_duration", ""))

    _add_h1(doc, "Team Structure")
    _add_table(
        doc,
        headers=["Role", "Count", "Responsibilities"],
        rows=[[t.get("role", ""), t.get("count", ""), t.get("responsibilities", "")]
              for t in structured_output.get("team_structure", [])],
        col_widths=[2.0, 0.8, 3.7],
    )

    _add_h1(doc, "Project Phases")
    for phase in structured_output.get("phases", []):
        _add_h2(doc, f"Phase {phase.get('phase_number', '')} — {phase.get('phase_name', '')}")
        _add_label_value(doc, "Duration", f"{phase.get('duration', '')}  |  Weeks {phase.get('start_week', '')}–{phase.get('end_week', '')}")
        _add_label_value(doc, "Objectives", "")
        for obj in phase.get("objectives", []):
            _add_bullet(doc, obj)
        tasks = phase.get("tasks", [])
        if tasks:
            _add_table(
                doc,
                headers=["Task ID", "Task Name", "Role", "Days", "Linked Requirements"],
                rows=[
                    [t.get("task_id", ""), t.get("task_name", ""), t.get("assignee_role", ""),
                     t.get("duration_days", ""), ", ".join(t.get("linked_requirements", []))]
                    for t in tasks
                ],
                col_widths=[0.8, 2.5, 1.5, 0.7, 1.0],
            )
        _add_label_value(doc, "Deliverables", "")
        for d in phase.get("deliverables", []):
            _add_bullet(doc, d)

    _add_h1(doc, "Risk Register")
    risks = structured_output.get("risk_register", [])
    if risks:
        _add_table(
            doc,
            headers=["ID", "Risk", "Probability", "Impact", "Mitigation", "Owner"],
            rows=[
                [r.get("risk_id", ""), r.get("description", ""), r.get("probability", ""),
                 r.get("impact", ""), r.get("mitigation", ""), r.get("owner", "")]
                for r in risks
            ],
            col_widths=[0.5, 2.0, 0.8, 0.8, 2.0, 0.5],
        )

    _add_h1(doc, "Definition of Done")
    for item in structured_output.get("definition_of_done", []):
        _add_bullet(doc, item)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Test Cases Exporter ─────────────────────────────────────────────────────

def export_testcases_to_docx(structured_output: dict) -> BytesIO:
    """Export Test Cases JSON to a fully ABHFL-branded Word document."""
    from datetime import date as dt

    doc = Document()
    _setup_page(doc)
    _add_logo_header(doc)
    _add_footer(doc)

    today = dt.today().strftime("%B %Y")
    _add_cover_page(doc, "Test Cases & Traceability Matrix", "Quality Assurance", date=today)

    summary = structured_output.get("test_summary", {})
    _add_h1(doc, "Test Summary")
    _add_label_value(doc, "Total Test Cases", str(summary.get("total_test_cases", 0)))
    _add_label_value(doc, "Coverage", summary.get("coverage_percentage", ""))
    cats = summary.get("test_categories", {})
    if cats:
        _add_table(
            doc,
            headers=["Category", "Count"],
            rows=[[k.title(), v] for k, v in cats.items()],
            col_widths=[3.0, 1.5],
        )

    _add_h1(doc, "Test Cases")
    for tc in structured_output.get("test_cases", []):
        _add_h2(doc, f"TC-{tc.get('sr_no', '')}  {tc.get('test_case_id', '')} — {tc.get('scenario_description', '')[:80]}")
        fields = [
            ("Module", tc.get("tc_module", "")),
            ("Sub-Module", tc.get("tc_sub_module", "")),
            ("Priority", tc.get("test_priority", "")),
            ("Classification", tc.get("test_classification", "")),
            ("Category", tc.get("test_category", "")),
            ("BRD Reference", tc.get("brd_fsd_reference", "")),
            ("Scenario ID", tc.get("scenario_id", "")),
            ("Importance", tc.get("importance", "")),
            ("Ownership", tc.get("ownership", "")),
            ("Release", tc.get("release", "")),
        ]
        _add_table(
            doc,
            headers=["Field", "Value"],
            rows=fields,
            col_widths=[2.0, 4.5],
        )
        _add_label_value(doc, "Path", tc.get("path", ""))
        _add_label_value(doc, "Pre-Requisite", tc.get("pre_requisite", ""))
        _add_label_value(doc, "Test Condition", tc.get("test_condition", ""))
        _add_label_value(doc, "Test Data", tc.get("test_data", ""))
        _add_label_value(doc, "Test Case Description", "")
        _add_body(doc, tc.get("test_case_description", ""))
        _add_label_value(doc, "Expected Result", tc.get("expected_result", ""))
        doc.add_paragraph()

    _add_h1(doc, "Requirements Traceability Matrix")
    matrix = structured_output.get("traceability_matrix", [])
    if matrix:
        _add_table(
            doc,
            headers=["Requirement ID", "Requirement Title", "Linked Test Cases", "Coverage Status"],
            rows=[
                [m.get("requirement_id", ""), m.get("requirement_title", ""),
                 ", ".join(m.get("linked_test_cases", [])), m.get("coverage_status", "")]
                for m in matrix
            ],
            col_widths=[1.2, 2.5, 1.8, 1.0],
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Effort Estimation Exporter ───────────────────────────────────────────────

def export_effort_to_docx(structured_output: dict) -> BytesIO:
    """Export Effort Estimation JSON to a fully ABHFL-branded Word document."""
    from datetime import date as dt

    doc = Document()
    _setup_page(doc)
    _add_logo_header(doc)
    _add_footer(doc)

    today = dt.today().strftime("%B %Y")
    _add_cover_page(doc, "Effort Estimation Report", structured_output.get("project_name", "Project"), date=today)

    summary = structured_output.get("estimation_summary", {})
    _add_h1(doc, "Estimation Summary")
    _add_label_value(doc, "Total Hours", str(summary.get("total_hours", 0)))
    _add_label_value(doc, "Duration", f"{summary.get('total_weeks', 0)} weeks / {summary.get('total_months', 0)} months")
    _add_label_value(doc, "Team Size", str(summary.get("team_size_recommended", 0)))
    _add_label_value(doc, "Confidence Level", summary.get("confidence_level", ""))
    _add_label_value(doc, "Methodology", summary.get("estimation_methodology", ""))

    _add_h1(doc, "Effort by Phase")
    _add_table(
        doc,
        headers=["Phase", "Hours", "% of Total"],
        rows=[[p.get("phase", ""), p.get("hours", ""), p.get("percentage_of_total", "")]
              for p in structured_output.get("by_phase", [])],
        col_widths=[3.0, 1.5, 1.5],
    )

    _add_h1(doc, "Effort by Role")
    _add_table(
        doc,
        headers=["Role", "Hours", "Daily Rate (INR)", "Total Cost (INR)"],
        rows=[[r.get("role", ""), r.get("hours", ""), r.get("daily_rate_usd", ""), r.get("total_cost_usd", "")]
              for r in structured_output.get("by_role", [])],
        col_widths=[2.5, 1.0, 1.5, 1.5],
    )

    _add_h1(doc, "Effort by Feature")
    _add_table(
        doc,
        headers=["Feature Area", "Linked Requirements", "Complexity", "Hours", "Notes"],
        rows=[
            [f.get("feature_area", ""), ", ".join(f.get("linked_requirements", [])),
             f.get("complexity", ""), f.get("hours", ""), f.get("notes", "")]
            for f in structured_output.get("by_feature", [])
        ],
        col_widths=[1.8, 1.5, 1.0, 0.7, 1.5],
    )

    _add_h1(doc, "Cost Estimate")
    cost = structured_output.get("cost_estimate", {})
    currency = cost.get("currency", "INR")
    _add_label_value(doc, "Currency", currency)
    _add_label_value(doc, "Low Estimate",  f"{currency} {cost.get('low_estimate', 0):,}")
    _add_label_value(doc, "Mid Estimate",  f"{currency} {cost.get('mid_estimate', 0):,}")
    _add_label_value(doc, "High Estimate", f"{currency} {cost.get('high_estimate', 0):,}")
    _add_body(doc, cost.get("notes", ""))

    _add_h1(doc, "Risks Affecting Estimate")
    risks = structured_output.get("risks_affecting_estimate", [])
    if risks:
        _add_table(
            doc,
            headers=["Risk", "Potential Impact (hours)", "Mitigation"],
            rows=[[r.get("risk", ""), r.get("potential_impact_hours", ""), r.get("mitigation", "")]
                  for r in risks],
            col_widths=[2.5, 1.5, 2.5],
        )

    _add_h1(doc, "Assumptions")
    for assumption in structured_output.get("assumptions", []):
        _add_bullet(doc, assumption)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
