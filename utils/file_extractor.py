"""
File Text Extractor Utility.

Extracts plain text from uploaded files:
  - PDF  → uses pdfplumber (primary) or PyPDF2 (fallback)
  - DOCX → uses python-docx
  - TXT  → native read with encoding detection

If extraction fails for any reason, returns None and the caller
should tell the user to paste the text manually instead.
"""

import logging
import os

from utils.image_extractor import extract_and_describe_images

logger = logging.getLogger(__name__)



def extract_text_from_file(file_path: str) -> str | None:
    """
    Extract plain text from a file at the given path.

    Args:
        file_path: Absolute path to the uploaded file.

    Returns:
        str: Extracted text, or None if extraction failed.
    """
    if not os.path.exists(file_path):
        logger.error(f'File not found: {file_path}')
        return None

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.pdf':
            text = _extract_from_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            text = _extract_from_docx(file_path)
        elif ext == '.txt':
            text = _extract_from_txt(file_path)
        elif ext in ('.xlsx', '.xls'):
            text = _extract_from_xlsx(file_path)
        elif ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp'):
            # Direct image upload — describe with AI vision via image_extractor
            text = ''
        else:
            logger.warning(f'Unsupported file type: {ext}')
            return None
    except Exception as e:
        logger.error(f'Text extraction failed for {file_path}: {e}')
        return None

    # ── Append embedded image descriptions (PDF / DOCX / direct image) ────────────
    if ext in ('.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp'):
        try:
            image_descriptions = extract_and_describe_images(file_path)
            if image_descriptions:
                text = (text or '')
                text += (
                    '\n\n'
                    '=== EMBEDDED IMAGE ANALYSIS ===\n'
                    '(Images described by AI vision for BRD context)\n\n'
                    + image_descriptions
                )
                logger.info(
                    f'Image analysis appended for {os.path.basename(file_path)} '
                    f'({len(image_descriptions)} chars)'
                )
        except Exception as img_exc:
            logger.warning(f'Image extraction skipped for {file_path}: {img_exc}')

    return text.strip() if text else None


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber, fallback to PyPDF2."""
    # Primary: pdfplumber (more accurate, handles complex layouts)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = '\n'.join(text_parts).strip()
        if text:
            logger.info(f'PDF extracted via pdfplumber: {len(text)} chars')
            return text
    except Exception as e:
        logger.warning(f'pdfplumber failed: {e} — trying PyPDF2')

    # Fallback: PyPDF2
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = '\n'.join(text_parts).strip()
        if text:
            logger.info(f'PDF extracted via PyPDF2: {len(text)} chars')
            return text
    except Exception as e:
        logger.error(f'PyPDF2 also failed: {e}')

    raise RuntimeError(f'Could not extract text from PDF: {file_path}')


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    text = '\n'.join(text_parts).strip()
    logger.info(f'DOCX extracted: {len(text)} chars')
    return text


def _extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file with encoding detection."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read().strip()
            logger.info(f'TXT extracted with {encoding}: {len(text)} chars')
            return text
        except UnicodeDecodeError:
            continue

    raise RuntimeError(f'Could not decode TXT file with any supported encoding: {file_path}')


def _extract_from_xlsx(file_path: str) -> str:
    """Extract text from XLSX using openpyxl, traversing all sheets."""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl package not installed. Cannot extract Excel.")

    # data_only=True ensures we get formula results rather than the formula text
    wb = openpyxl.load_workbook(file_path, data_only=True)
    text_parts = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")
        for row in sheet.iter_rows(values_only=True):
            # filter out None and empty strings, convert others to string
            row_texts = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if row_texts:
                text_parts.append(" | ".join(row_texts))
                
    text = '\n'.join(text_parts).strip()
    if text:
        logger.info(f'XLSX extracted: {len(text)} chars')
        return text
    else:
        logger.warning(f'XLSX {file_path} was empty or unreadable.')
        return ""
